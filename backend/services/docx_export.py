"""
DOCX Export Service
Generates Word documents for tests and study materials
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import html
from html.parser import HTMLParser


class HTMLToTextParser(HTMLParser):
    """Simple HTML to plain text parser"""
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []

    def handle_data(self, d):
        self.text.append(d)

    def get_data(self):
        return ''.join(self.text)


def strip_html(html_text):
    """Strip HTML tags from text"""
    parser = HTMLToTextParser()
    parser.feed(html_text)
    return parser.get_data()


def generate_test_docx(test_data):
    """
    Generate DOCX for a test (student version without answers)

    Args:
        test_data: Test object with assignments and questions

    Returns:
        BytesIO: DOCX file content
    """
    doc = Document()


    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading(test_data['title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)

    from datetime import datetime
    date_str = datetime.fromisoformat(test_data['created_at'].replace('Z', '+00:00')).strftime('%B %d, %Y')
    date_para = doc.add_paragraph(f"Izveidots: {date_str}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    instructions = doc.add_paragraph()
    instructions.add_run("Instrukcija: ").bold = True
    instructions.add_run("Atbildiet uz visiem jautājumiem. Rakstiet savas atbildes skaidri.")
    instructions.runs[0].font.size = Pt(10)
    instructions.runs[1].font.size = Pt(10)
    doc.add_paragraph()

    for assignment_idx, assignment in enumerate(test_data['assignments']):
        assignment_num = assignment_idx + 1

        assignment_heading = doc.add_heading(f"Uzdevums {assignment_num}: {assignment['title']}", level=1)
        assignment_heading.runs[0].font.size = Pt(13)

        # Skip assignment description (usually contains AI generation prompts)
        # if assignment.get('description'):
        #     desc_para = doc.add_paragraph(assignment['description'])
        #     desc_para.runs[0].font.size = Pt(10)

        max_points = doc.add_paragraph()
        max_points_run = max_points.add_run(f"Maksimālais punktu skaits: {assignment['max_points']}")
        max_points_run.italic = True
        max_points_run.font.size = Pt(10)

        doc.add_paragraph()

        for question_idx, question in enumerate(assignment['questions']):
            question_num = question_idx + 1

            q_header = doc.add_paragraph()
            q_header_run = q_header.add_run(f"Jautājums {question_num}")
            q_header_run.bold = True
            q_header_run.font.size = Pt(10)

            q_header.add_run(f" ({question['points']} punkti)")
            q_header.runs[1].font.size = Pt(10)

            q_text = html.unescape(question['question_text'])
            q_para = doc.add_paragraph(q_text)
            q_para.runs[0].font.size = Pt(10)

            if question.get('options') and len(question['options']) > 0:
                if question['question_type'] == 'matching':
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    num_rows = len(question['options']) + 1
                    table = doc.add_table(rows=num_rows, cols=3)
                    table.style = 'Table Grid'

                    for row in table.rows:
                        row.cells[0].width = Inches(2.2)
                        row.cells[1].width = Inches(1.2)
                        row.cells[2].width = Inches(2.2)

                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Kreisā puse'
                    header_cells[1].text = ''
                    header_cells[2].text = 'Labā puse'

                    for cell in header_cells:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'D9D9D9')
                        cell._element.get_or_add_tcPr().append(shading_elm)

                    for opt_idx, option in enumerate(question['options']):
                        row = table.rows[opt_idx + 1]
                        option_text = option['option_text']
                        parts = option_text.split('|')
                        left_part = html.unescape(parts[0]) if len(parts) > 0 else ''
                        right_part = html.unescape(parts[1]) if len(parts) > 1 else ''

                        left_num = opt_idx + 1
                        right_num = opt_idx + 1

                        left_cell = row.cells[0]
                        left_cell.text = f'{left_num}. {left_part}'

                        row.cells[1].text = ''

                        right_cell = row.cells[2]
                        right_cell.text = f'{right_num}. {right_part}'

                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                    instruction_para = doc.add_paragraph()
                    instruction_run = instruction_para.add_run(
                        'Velciet līnijas, lai saskaņotu elementus no kreisās kolonnas ar labo kolonnu.'
                    )
                    instruction_run.italic = True
                    instruction_run.font.size = Pt(9)

                else:
                    for opt_idx, option in enumerate(question['options']):
                        option_letter = chr(65 + opt_idx)
                        option_text = html.unescape(option['option_text'])

                        opt_para = doc.add_paragraph(style='List Bullet')
                        opt_run = opt_para.add_run(f"{option_letter}. {option_text}")
                        opt_run.font.size = Pt(10)

            if question['question_type'] in ['short_answer', 'fill_in_blank']:
                doc.add_paragraph("_" * 80)
            elif question['question_type'] == 'long_answer':
                for _ in range(4):
                    doc.add_paragraph("_" * 80)

            doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_study_material_docx(material_data):
    """
    Generate DOCX for study material

    Args:
        material_data: Study material object with summary and terms

    Returns:
        BytesIO: DOCX file content
    """
    doc = Document()


    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading(material_data['title'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)

    from datetime import datetime
    date_str = datetime.fromisoformat(material_data['created_at'].replace('Z', '+00:00')).strftime('%B %d, %Y')
    date_para = doc.add_paragraph(f"Izveidots: {date_str}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    summary_heading = doc.add_heading("Kopsavilkums", level=1)
    summary_heading.runs[0].font.size = Pt(13)

    # Handle HTML content in summary (strip HTML tags)
    summary_content = material_data['content'].get('summary', '')
    summary_text = strip_html(summary_content)

    summary_para = doc.add_paragraph(summary_text)
    summary_para.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    terms_heading = doc.add_heading("Galvenie termini", level=1)
    terms_heading.runs[0].font.size = Pt(13)

    terms = material_data['content'].get('terms', [])
    for term in terms:
        term_para = doc.add_paragraph()
        term_run = term_para.add_run(term['name'])
        term_run.bold = True
        term_run.font.size = Pt(12)

        definition = html.unescape(term['definition'])
        def_para = doc.add_paragraph(definition)
        def_para.runs[0].font.size = Pt(11)

        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
